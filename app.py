import streamlit as st
import streamlit.components.v1 as components
import os
import io
import sqlite3
import hashlib
import binascii
import base64
import tempfile
from datetime import datetime
from dotenv import load_dotenv
from groq import Groq
from docx import Document
from pypdf import PdfReader

# Page setup — sidebar is now the main navigation, so keep it open by default
st.set_page_config(
    page_title="Teacher Assistant",
    page_icon="🧑‍🏫",
    layout="centered",
    initial_sidebar_state="expanded"
)

# Load your API key from .env (for local use) or Streamlit secrets (for cloud deployment)
load_dotenv()
groq_api_key = os.getenv("GROQ_API_KEY") or st.secrets.get("GROQ_API_KEY", None)
client = Groq(api_key=groq_api_key, timeout=60.0)

# A secret code shared only with teachers — required to create an account, so students can't sign up
TEACHER_ACCESS_CODE = os.getenv("TEACHER_ACCESS_CODE") or st.secrets.get("TEACHER_ACCESS_CODE", None)

# ---------- Helper: time-based greeting ----------
def get_greeting():
    hour = datetime.now().hour
    if hour < 12:
        return "Good morning"
    elif hour < 18:
        return "Good afternoon"
    else:
        return "Good evening"

# ---------- General UI polish: chat bubbles, spacing, rounded controls ----------
st.markdown("""
<style>
/* Chat bubble styling — different tint for user vs assistant */
[data-testid="stChatMessage"] {
    border-radius: 16px;
    padding: 10px 16px;
    margin-bottom: 10px;
}
[data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarUser"]) {
    background-color: rgba(46, 125, 107, 0.08);
}
[data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarAssistant"]) {
    background-color: rgba(0, 0, 0, 0.035);
}

/* Fix avatar sizing — give them a proper fixed circle instead of a clipped default */
[data-testid="stChatMessageAvatarUser"],
[data-testid="stChatMessageAvatarAssistant"] {
    width: 38px !important;
    height: 38px !important;
    min-width: 38px !important;
    border-radius: 50% !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    font-size: 20px !important;
    line-height: 1 !important;
    overflow: hidden;
}
[data-testid="stChatMessageAvatarUser"] {
    background-color: rgba(46, 125, 107, 0.18) !important;
}
[data-testid="stChatMessageAvatarAssistant"] {
    background-color: rgba(0, 0, 0, 0.08) !important;
}

/* Sidebar button spacing — slightly tighter, consistent rounding */
section[data-testid="stSidebar"] button {
    border-radius: 10px !important;
    margin-bottom: 2px;
}

/* General rounded corners on form containers and expanders */
div[data-testid="stExpander"] {
    border-radius: 14px;
}

/* Reduce default top padding so the greeting sits closer to the top */
.main .block-container {
    padding-top: 2rem;
}
</style>
""", unsafe_allow_html=True)

# ---------- Soft green glow background (decorative, sits behind everything) ----------
st.markdown("""
<div class="glow-bg">
    <div class="glow-orb glow-orb-1"></div>
    <div class="glow-orb glow-orb-2"></div>
    <div class="glow-orb glow-orb-3"></div>
</div>
<style>
.glow-bg {
    position: fixed;
    top: 0; left: 0; right: 0; bottom: 0;
    z-index: -1;
    overflow: hidden;
    pointer-events: none;
}
.glow-orb {
    position: absolute;
    border-radius: 50%;
    filter: blur(70px);
    opacity: 0.65;
}
.glow-orb-1 {
    width: 420px; height: 420px;
    background: #2E7D6B;
    top: -120px; left: -100px;
}
.glow-orb-2 {
    width: 380px; height: 380px;
    background: #6FBF9B;
    bottom: -100px; right: -80px;
}
.glow-orb-3 {
    width: 320px; height: 320px;
    background: #A8D5C0;
    top: 35%; left: 55%;
}
</style>
""", unsafe_allow_html=True)

# ---------- Database setup (saves chats to a local file, organized into separate conversations) ----------
DB_PATH = "chat_history.db"

def init_db():
    conn = sqlite3.connect(DB_PATH)

    # Users table (accounts)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL UNIQUE,
            password_hash TEXT NOT NULL,
            salt TEXT NOT NULL,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    conn.execute("""
        CREATE TABLE IF NOT EXISTS feedback (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            message TEXT NOT NULL,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    conn.execute("""
        CREATE TABLE IF NOT EXISTS conversations (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            title TEXT NOT NULL DEFAULT 'New Chat',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    # Migration: add user_id column if this conversations table pre-dates login
    convo_cols = [row[1] for row in conn.execute("PRAGMA table_info(conversations)")]
    if "user_id" not in convo_cols:
        conn.execute("ALTER TABLE conversations ADD COLUMN user_id INTEGER")

    # Check if an old-style messages table (no conversation_id) already exists
    existing_cols = [row[1] for row in conn.execute("PRAGMA table_info(messages)")]

    if existing_cols and "conversation_id" not in existing_cols:
        # Old schema found — migrate it into a single conversation instead of losing history
        conn.execute("ALTER TABLE messages RENAME TO messages_old")
        conn.execute("""
            CREATE TABLE messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                conversation_id INTEGER NOT NULL,
                role TEXT NOT NULL,
                content TEXT NOT NULL,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (conversation_id) REFERENCES conversations(id)
            )
        """)
        cur = conn.execute("INSERT INTO conversations (title) VALUES ('Previous Chat')")
        migrated_convo_id = cur.lastrowid
        old_rows = conn.execute("SELECT role, content FROM messages_old ORDER BY id ASC").fetchall()
        for role, content in old_rows:
            conn.execute(
                "INSERT INTO messages (conversation_id, role, content) VALUES (?, ?, ?)",
                (migrated_convo_id, role, content)
            )
        conn.execute("DROP TABLE messages_old")
    else:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS messages (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                conversation_id INTEGER NOT NULL,
                role TEXT NOT NULL,
                content TEXT NOT NULL,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (conversation_id) REFERENCES conversations(id)
            )
        """)

    conn.commit()
    conn.close()

init_db()

# ---------- Password hashing (never store plain-text passwords) ----------
def hash_password(password, salt=None):
    if salt is None:
        salt = os.urandom(16)
    dk = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, 100000)
    return binascii.hexlify(dk).decode(), binascii.hexlify(salt).decode()

def verify_password(password, stored_hash, stored_salt):
    salt = binascii.unhexlify(stored_salt)
    dk = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, 100000)
    return binascii.hexlify(dk).decode() == stored_hash

# ---------- User account functions ----------
def create_user(username, password):
    conn = sqlite3.connect(DB_PATH)
    existing = conn.execute("SELECT id FROM users WHERE username = ?", (username,)).fetchone()
    if existing:
        conn.close()
        return None  # username already taken
    pw_hash, salt = hash_password(password)
    cur = conn.execute(
        "INSERT INTO users (username, password_hash, salt) VALUES (?, ?, ?)",
        (username, pw_hash, salt)
    )
    conn.commit()
    new_user_id = cur.lastrowid
    # One-time: claim any old conversations that existed before accounts were added
    conn.execute("UPDATE conversations SET user_id = ? WHERE user_id IS NULL", (new_user_id,))
    conn.commit()
    conn.close()
    return new_user_id

def authenticate_user(username, password):
    conn = sqlite3.connect(DB_PATH)
    row = conn.execute(
        "SELECT id, password_hash, salt FROM users WHERE username = ?", (username,)
    ).fetchone()
    conn.close()
    if not row:
        return None
    user_id, pw_hash, salt = row
    if verify_password(password, pw_hash, salt):
        return user_id
    return None

def save_feedback(user_id, message):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("INSERT INTO feedback (user_id, message) VALUES (?, ?)", (user_id, message))
    conn.commit()
    conn.close()

def list_conversations(user_id):
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute(
        "SELECT id, title FROM conversations WHERE user_id = ? ORDER BY id DESC", (user_id,)
    ).fetchall()
    conn.close()
    return [{"id": r[0], "title": r[1]} for r in rows]

def create_conversation(user_id, title="New Chat"):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.execute("INSERT INTO conversations (user_id, title) VALUES (?, ?)", (user_id, title))
    conn.commit()
    new_id = cur.lastrowid
    conn.close()
    return new_id

def rename_conversation(conversation_id, title):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("UPDATE conversations SET title = ? WHERE id = ?", (title, conversation_id))
    conn.commit()
    conn.close()

def delete_conversation(conversation_id):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("DELETE FROM messages WHERE conversation_id = ?", (conversation_id,))
    conn.execute("DELETE FROM conversations WHERE id = ?", (conversation_id,))
    conn.commit()
    conn.close()

def load_messages(conversation_id):
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute(
        "SELECT id, role, content FROM messages WHERE conversation_id = ? ORDER BY id ASC",
        (conversation_id,)
    ).fetchall()
    conn.close()
    return [{"id": r[0], "role": r[1], "content": r[2]} for r in rows]

def list_all_assistant_messages(user_id):
    """All AI-generated documents across every one of this teacher's chats — newest first."""
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute("""
        SELECT messages.id, messages.content
        FROM messages
        JOIN conversations ON messages.conversation_id = conversations.id
        WHERE conversations.user_id = ? AND messages.role = 'assistant'
        ORDER BY messages.id DESC
    """, (user_id,)).fetchall()
    conn.close()
    return [{"id": r[0], "content": r[1]} for r in rows]

def save_message(conversation_id, role, content):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.execute(
        "INSERT INTO messages (conversation_id, role, content) VALUES (?, ?, ?)",
        (conversation_id, role, content)
    )
    conn.commit()
    new_id = cur.lastrowid
    conn.close()
    return new_id

def update_message(msg_id, content):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("UPDATE messages SET content = ? WHERE id = ?", (content, msg_id))
    conn.commit()
    conn.close()

# ---------- Login / Signup gate ----------
if "user_id" not in st.session_state:
    st.session_state.user_id = None
    st.session_state.username = None

if st.session_state.user_id is None:
    st.title("🧑‍🏫 Teacher Assistant")
    st.caption("Sign in to access your lesson plans, quizzes, and chat history")

    login_tab, signup_tab = st.tabs(["Log In", "Sign Up"])

    with login_tab:
        with st.form("login_form"):
            login_username = st.text_input("Username")
            login_password = st.text_input("Password", type="password")
            login_submitted = st.form_submit_button("Log In")

            if login_submitted:
                if not login_username or not login_password:
                    st.warning("Please enter both a username and password.")
                else:
                    result_user_id = authenticate_user(login_username, login_password)
                    if result_user_id:
                        st.session_state.user_id = result_user_id
                        st.session_state.username = login_username
                        st.rerun()
                    else:
                        st.error("Incorrect username or password.")

    with signup_tab:
        with st.form("signup_form"):
            st.caption("Ask your school admin for the teacher access code if you don't have one.")
            new_username = st.text_input("Choose a username")
            new_password = st.text_input("Choose a password", type="password")
            confirm_password = st.text_input("Confirm password", type="password")
            entered_access_code = st.text_input("Teacher access code", type="password")
            signup_submitted = st.form_submit_button("Create Account")

            if signup_submitted:
                if not new_username or not new_password or not entered_access_code:
                    st.warning("Please fill in all fields.")
                elif not TEACHER_ACCESS_CODE:
                    st.error("Sign-up isn't configured yet — the site admin needs to set a teacher access code.")
                elif entered_access_code != TEACHER_ACCESS_CODE:
                    st.error("That access code isn't correct. Ask your school admin for the right one.")
                elif new_password != confirm_password:
                    st.warning("Passwords don't match.")
                elif len(new_password) < 6:
                    st.warning("Password should be at least 6 characters.")
                else:
                    created_id = create_user(new_username, new_password)
                    if created_id:
                        st.session_state.user_id = created_id
                        st.session_state.username = new_username
                        st.rerun()
                    else:
                        st.error("That username is already taken. Try another.")

    st.stop()  # don't render the rest of the app until logged in

# ---------- Session state setup ----------
if "current_conversation_id" not in st.session_state:
    existing = list_conversations(st.session_state.user_id)
    if existing:
        st.session_state.current_conversation_id = existing[0]["id"]  # most recent chat
    else:
        st.session_state.current_conversation_id = create_conversation(st.session_state.user_id)

if "messages" not in st.session_state:
    st.session_state.messages = load_messages(st.session_state.current_conversation_id)

if "mode" not in st.session_state:
    st.session_state.mode = None  # which form is currently open

# ---------- Helper: a button that reads text aloud using the browser's built-in voice ----------
def generate_speech_audio(text, voice="autumn"):
    """Calls Groq's Orpheus text-to-speech model and returns real audio bytes (WAV).
    (Groq retired the older PlayAI models in favor of Orpheus.)"""
    response = client.audio.speech.create(
        model="canopylabs/orpheus-v1-english",
        voice=voice,
        input=text[:2000],  # keep requests reasonably short
        response_format="wav"
    )
    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
        response.write_to_file(tmp.name)
        tmp_path = tmp.name
    with open(tmp_path, "rb") as f:
        audio_bytes = f.read()
    os.remove(tmp_path)
    return audio_bytes

def speak_button(text, key):
    cache_key = f"tts_audio_{key}"
    if st.button("🔊 Listen", key=f"listen_btn_{key}", use_container_width=True):
        if cache_key not in st.session_state:
            with st.spinner("Generating voice..."):
                try:
                    voice_choice = st.session_state.get("tts_voice", "autumn")
                    st.session_state[cache_key] = generate_speech_audio(text, voice_choice)
                except Exception as e:
                    st.error(f"Couldn't generate voice: {e}")
                    return
        audio_bytes = st.session_state[cache_key]
        b64 = base64.b64encode(audio_bytes).decode()
        html_code = f"""
        <audio autoplay controls style="width:100%; margin-top:6px;">
            <source src="data:audio/wav;base64,{b64}" type="audio/wav">
        </audio>
        """
        components.html(html_code, height=60)

# ---------- Helper: turn AI text into a downloadable Word doc ----------
def text_to_docx_bytes(text):
    doc = Document()
    lines = text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ---- Detect a markdown table: a "| ... |" row followed by a "|---|---|" separator row ----
        if stripped.startswith("|") and stripped.endswith("|") and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            is_separator = next_line.startswith("|") and set(next_line.replace("|", "").replace(":", "").strip()) <= {"-", " "}

            if is_separator:
                # Collect every consecutive table row
                table_rows = [stripped]
                j = i + 2  # skip the separator row itself
                while j < len(lines) and lines[j].strip().startswith("|") and lines[j].strip().endswith("|"):
                    table_rows.append(lines[j].strip())
                    j += 1

                parsed_rows = [
                    [cell.strip() for cell in row.strip("|").split("|")]
                    for row in table_rows
                ]
                num_cols = max(len(r) for r in parsed_rows)

                table = doc.add_table(rows=0, cols=num_cols)
                table.style = "Table Grid"

                for row_idx, row_values in enumerate(parsed_rows):
                    row_cells = table.add_row().cells
                    for col_idx in range(num_cols):
                        cell_text = row_values[col_idx] if col_idx < len(row_values) else ""
                        row_cells[col_idx].text = cell_text
                        if row_idx == 0:  # bold the header row
                            for para in row_cells[col_idx].paragraphs:
                                for run in para.runs:
                                    run.bold = True

                doc.add_paragraph("")  # small spacing gap after the table
                i = j
                continue

        # ---- Not a table row — handle normally ----
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            p = doc.add_paragraph()
            p.add_run(stripped[2:-2]).bold = True
        else:
            doc.add_paragraph(line)

        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------- Helper: read text out of an uploaded file ----------
def read_uploaded_file(uploaded_file):
    if uploaded_file.name.endswith(".docx"):
        doc = Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs)
    elif uploaded_file.name.endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    else:  # assume plain text
        return uploaded_file.read().decode("utf-8", errors="ignore")

# ---------- Sidebar: main navigation (Claude-style left panel) ----------
with st.sidebar:
    st.markdown("### 🧑‍🏫 Teacher Assistant")
    st.caption(f"👤 Logged in as **{st.session_state.username}**")

    if "dark_mode" not in st.session_state:
        st.session_state.dark_mode = False
    # A plain, direct toggle — applies immediately on click, no popover involved
    st.session_state.dark_mode = st.toggle("🌙 Dark mode", value=st.session_state.dark_mode, key="main_dark_toggle")

    if "tts_voice" not in st.session_state:
        st.session_state.tts_voice = "autumn"
    voice_options = ["autumn", "diana", "hannah", "austin", "daniel", "troy"]
    st.session_state.tts_voice = st.selectbox(
        "🔊 AI voice", voice_options,
        index=voice_options.index(st.session_state.tts_voice),
        key="sidebar_voice_select"
    )

    with st.popover("🚩 Send feedback", use_container_width=True):
        fb_text = st.text_area("What's on your mind?", key="feedback_text")
        if st.button("Submit", key="feedback_submit"):
            if fb_text.strip():
                save_feedback(st.session_state.user_id, fb_text.strip())
                st.success("Thanks! Your feedback was saved.")
            else:
                st.warning("Type something first.")

    # Real, working nav — shown in both light and dark mode
    if st.button("🧭 Explore", key="nav_explore", use_container_width=True):
        st.session_state.show_explore = not st.session_state.get("show_explore", False)
    if st.session_state.get("show_explore"):
        st.caption("Quick things to try:")
        for label, prompt_text in [
            ("📝 Make a lesson plan", "Create a lesson plan for a topic I'll describe"),
            ("❓ Build a quick quiz", "Create a short quiz on a topic I'll describe"),
            ("💡 Brainstorm", "Help me brainstorm creative teaching ideas for a topic I'll describe"),
            ("💻 Code", "Help me write or fix a spreadsheet formula, script, or piece of code for a classroom task I'll describe."),
            ("🎓 Get Advice", "Give me practical teaching advice on a classroom challenge I'll describe."),
        ]:
            if st.button(label, key=f"explore_{label}", use_container_width=True):
                st.session_state.pending_voice_prompt = prompt_text
                st.session_state.show_explore = False
                st.rerun()
        if st.button("🌐 Web search", key="explore_websearch", use_container_width=True, help="Needs a separate search API — not connected yet"):
            st.toast("🌐 Web search needs a separate search service to be connected first.")

    if st.button("🗂️ Categories", key="nav_categories", use_container_width=True):
        st.session_state.show_categories = not st.session_state.get("show_categories", False)
    if st.session_state.get("show_categories"):
        st.caption("📝 **Lesson Plan** — structured plans with objectives & stages")
        st.caption("❓ **Quiz/Test** — sectioned exams with mark allocations")
        st.caption("📋 **Assignment** — scenario-based tasks")
        st.caption("💬 **Report Comment** — student report card comments")
        st.caption("📄 **Summarize Doc** — condense an uploaded file")
        st.caption("📓 **Generate Notes** — study notes on a topic")

    if st.button("📚 Library", key="nav_library", use_container_width=True):
        st.session_state.show_library = not st.session_state.get("show_library", False)
    if st.session_state.get("show_library"):
        library_docs = list_all_assistant_messages(st.session_state.user_id)
        if not library_docs:
            st.caption("Nothing generated yet — your documents will appear here.")
        else:
            for doc_row in library_docs[:15]:
                doc_preview = doc_row["content"].strip().split("\n")[0][:45]
                st.caption(f"📄 {doc_preview}...")

    if st.session_state.dark_mode:
        st.markdown("""
        <style>
        .stApp {
            background-color: #0F0F0F;
            color: #EAEAEA;
        }
        section[data-testid="stSidebar"] {
            background-color: #161616;
        }
        section[data-testid="stSidebar"] * {
            color: #EAEAEA !important;
        }
        /* User bubble: dark gray card, right-aligned like the reference */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarUser"]) {
            background-color: #262626 !important;
            flex-direction: row-reverse;
            margin-left: auto;
            max-width: 75%;
            border-radius: 18px;
        }
        /* Assistant message: no bubble, just plain text on the dark background */
        [data-testid="stChatMessage"]:has([data-testid="stChatMessageAvatarAssistant"]) {
            background-color: transparent !important;
            max-width: 100%;
        }
        .st-key-input_bar_wrapper {
            background:
                linear-gradient(rgba(28,28,28,0.92), rgba(28,28,28,0.92)) padding-box,
                linear-gradient(90deg, #EA4335, #4285F4, #34A853, #FBBC05, #EA4335) border-box !important;
            background-size: 100% 100%, 300% 100% !important;
        }
        div[data-testid="stTextInput"] input,
        div[data-testid="stTextArea"] textarea,
        div[data-testid="stSelectbox"] div,
        div[data-testid="stNumberInput"] input {
            background-color: #1C1C1C !important;
            color: #EAEAEA !important;
        }
        /* Buttons everywhere — sidebar, forms, download */
        button, .stDownloadButton button {
            background-color: #1C1C1C !important;
            color: #EAEAEA !important;
            border: 1px solid #3A3A3A !important;
            box-shadow: 0 1px 3px rgba(0,0,0,0.3);
        }
        button:hover, .stDownloadButton button:hover {
            border-color: #8FE388 !important;
            color: #8FE388 !important;
        }
        /* Consistent breathing room between stacked buttons */
        div[data-testid="stVerticalBlock"] > div:has(> div > button) {
            margin-bottom: 2px;
        }
        /* New Chat button — bright accent, like the reference's green pill */
        .st-key-new_chat_btn button {
            background-color: #8FE388 !important;
            color: #0F0F0F !important;
            border: none !important;
            font-weight: 600;
        }
        .st-key-new_chat_btn button:hover {
            background-color: #A6EE9F !important;
            color: #0F0F0F !important;
        }
        /* Forms, expanders, containers */
        div[data-testid="stForm"],
        div[data-testid="stExpander"] {
            background-color: #171717 !important;
            border-color: #333 !important;
        }
        /* Dark, rounded code blocks */
        pre, code {
            background-color: #1A1A1A !important;
            color: #D4D4D4 !important;
            border-radius: 10px !important;
        }
        pre {
            padding: 12px !important;
            border: 1px solid #2E2E2E !important;
        }
        /* Dividers */
        hr {
            border-color: #333 !important;
        }
        /* Captions and helper text */
        [data-testid="stCaptionContainer"], .stCaption, small {
            color: #999999 !important;
        }
        /* Links */
        a {
            color: #8FE388 !important;
        }
        /* Placeholder text in inputs */
        ::placeholder {
            color: #888888 !important;
            opacity: 1;
        }
        /* Tabs (used on the login screen) */
        button[data-baseweb="tab"] {
            color: #999999 !important;
        }
        button[data-baseweb="tab"][aria-selected="true"] {
            color: #8FE388 !important;
            border-color: #8FE388 !important;
        }
        .glow-orb {
            opacity: 0.5 !important;
        }
        [data-testid="stChatMessageAvatarUser"] {
            background-color: rgba(143, 227, 136, 0.25) !important;
        }
        [data-testid="stChatMessageAvatarAssistant"] {
            background-color: rgba(255, 255, 255, 0.10) !important;
        }
        </style>
        """, unsafe_allow_html=True)

    if st.button("🚪 Log Out", use_container_width=True):
        for key in ["user_id", "username", "current_conversation_id", "messages", "mode", "teacher_name"]:
            if key in st.session_state:
                del st.session_state[key]
        st.rerun()

    st.divider()
    st.caption("QUICK ACTIONS")

    if st.button("📝 Lesson Plan", use_container_width=True):
        st.session_state.mode = "lesson_plan"
    if st.button("❓ Quiz / Test", use_container_width=True):
        st.session_state.mode = "quiz"
    if st.button("📋 Assignment", use_container_width=True):
        st.session_state.mode = "assignment"
    if st.button("💬 Report Comment", use_container_width=True):
        st.session_state.mode = "comment"
    if st.button("📄 Summarize Doc", use_container_width=True):
        st.session_state.mode = "summarize"
    if st.button("📓 Generate Notes", use_container_width=True):
        st.session_state.mode = "notes"

    st.divider()
    if st.button("➕ New Chat", use_container_width=True, key="new_chat_btn"):
        new_id = create_conversation(st.session_state.user_id)
        st.session_state.current_conversation_id = new_id
        st.session_state.messages = []
        st.session_state.mode = None
        st.rerun()

    st.caption("YOUR CHATS")
    conversations = list_conversations(st.session_state.user_id)
    for convo in conversations:
        is_current = convo["id"] == st.session_state.current_conversation_id
        label = ("🟢 " if is_current else "") + convo["title"]
        col_chat, col_del = st.columns([4, 1])
        with col_chat:
            if st.button(label, key=f"convo_{convo['id']}", use_container_width=True):
                st.session_state.current_conversation_id = convo["id"]
                st.session_state.messages = load_messages(convo["id"])
                st.session_state.mode = None
                st.rerun()
        with col_del:
            if st.button("🗑️", key=f"del_{convo['id']}"):
                delete_conversation(convo["id"])
                if convo["id"] == st.session_state.current_conversation_id:
                    remaining = list_conversations(st.session_state.user_id)
                    if remaining:
                        st.session_state.current_conversation_id = remaining[0]["id"]
                        st.session_state.messages = load_messages(remaining[0]["id"])
                    else:
                        st.session_state.current_conversation_id = create_conversation(st.session_state.user_id)
                        st.session_state.messages = []
                st.rerun()

# ---------- Greeting + example prompts (shown at the top of the main area) ----------
greeting_name = f", {st.session_state.username}" if st.session_state.username else ""
st.title(f"{get_greeting()}{greeting_name} 👋")
st.caption("What would you like help with today?")

quick_prompt = None  # will hold the built prompt once a form is submitted


# ---------- Lesson Plan form ----------
if st.session_state.mode == "lesson_plan":
    with st.form("lesson_plan_form"):
        st.subheader("Lesson Plan Details")
        subject = st.text_input("Subject", placeholder="e.g. Design and Technology")
        topic = st.text_input("Topic / Sub-topic", placeholder="e.g. 2.1 Materials and Manufacturing - Heat Treatment")
        grade = st.text_input("Class / Grade", placeholder="e.g. Form 2")
        duration = st.selectbox("Duration", ["40 Minutes", "80 Minutes", "120 Minutes"])
        submitted = st.form_submit_button("Generate Lesson Plan")

        if submitted:
            quick_prompt = f"""Create a full lesson plan using this exact structure:

Header fields: Name of Teacher (leave blank), Date (leave blank), Class: {grade}, Duration: {duration}, Subject: {subject}, Topic: {topic}, Sub-topic, General Competence(s), Lesson Goal, Specific Competences, Rationale, Prior Knowledge, References, Learning Environment, Teaching and Learning Materials/Resources, Expected Standard.

Then a Lesson Progression table with 4 columns (Stages | Teacher's Role | Learners' Role | Assessment Criteria) with these stages:
- Introduction (10 mins): a short real-world scenario/question to hook learners, followed by discussion
- Lesson Development (50 mins): numbered teaching steps including a demonstration and a hands-on learner task
- Exercise/Assessment (15 mins): an individual task with clear criteria
- Homework: one short task or question
- Conclusion (5 mins): a recap and a review question

End with a Lesson Evaluation section (Challenges / Teacher's response) left blank for after the lesson.

Subject: {subject}
Topic: {topic}
Class: {grade}"""

# ---------- Quiz / Test form ----------
elif st.session_state.mode == "quiz":
    with st.form("quiz_form"):
        st.subheader("Quiz / Test Details")
        subject = st.text_input("Subject", placeholder="e.g. Chemistry")
        topics = st.text_input("Topics to include (comma-separated)", placeholder="e.g. Periodic table, Atomic structure")
        num_questions = st.number_input("Number of questions", min_value=1, max_value=40, value=10)
        question_types = st.multiselect(
            "Question types to include",
            ["Multiple Choice", "Structured / Short Answer", "Essay (choose 1 of a few)"],
            default=["Multiple Choice", "Structured / Short Answer"]
        )
        difficulty = st.selectbox("Difficulty", ["Easy", "Medium", "Hard", "Mixed"])
        submitted = st.form_submit_button("Generate Quiz / Test")

        if submitted:
            quick_prompt = f"""Create a test/quiz using this exact structure, matching a formal school exam paper:

Header: School name placeholder, Department, Subject: {subject}, Form/Class placeholder, Test name, Duration placeholder, Name field, Total marks.

Instructions block: number of sections, how to answer, any special instructions (e.g. use of drafting tools if relevant).

Then organize into sections based on the question types requested: {', '.join(question_types)}.
- Multiple Choice section: format each question exactly like this style (NOT a table):
  1) [question text]
  * A) [option]
  * B) [option]
  * C) [option]
  * D) [option]
  Each question worth 1 mark. Do NOT mark the correct answer inline — put all correct answers together in an Answer Key at the very end.
- Structured/Short Answer section: numbered questions with lettered sub-parts (a), (b), (c), each with a mark allocation in brackets like [2].
- Essay section (if included): present a few options and instruct the student to answer only one, worth more marks (e.g. [20]).

Include an Answer Key at the very end, separate from the questions.

Subject: {subject}
Topics to cover: {topics}
Number of questions: {num_questions}
Difficulty: {difficulty}"""

# ---------- Assignment form ----------
elif st.session_state.mode == "assignment":
    with st.form("assignment_form"):
        st.subheader("Assignment Details")
        subject = st.text_input("Subject", placeholder="e.g. Design and Technology")
        topics = st.text_input("Topics / Scenarios to cover (comma-separated)", placeholder="e.g. Heat treatment, Holding tools, Plastics")
        duration = st.text_input("Duration to complete", placeholder="e.g. 3 weeks")
        due_date = st.text_input("Due date", placeholder="e.g. 7th September, 2026")
        submitted = st.form_submit_button("Generate Assignment")

        if submitted:
            quick_prompt = f"""Create an assignment using this exact structure:

Header: School/Department placeholder, Subject: {subject}, Duration: {duration}, Due Date: {due_date}, Name field.

Instructions block: number of sections, answer-all instruction, any special notes (e.g. use of drafting tools for graphics sections).

Then organize into scenario-based tasks — for each topic listed, write:
- A short real-world Scenario paragraph setting up a practical problem
- A TASK section with numbered/lettered sub-questions, each with a mark allocation in brackets like [2]
- Include a mix of short-answer, listing, explaining, and (where relevant to the subject) a drawing/sketching task

If the subject involves drafting or graphics, add a final section instructing students to complete that part on A3 paper.

Subject: {subject}
Topics/Scenarios: {topics}"""

# ---------- Report Comment (simple, no form) ----------
elif st.session_state.mode == "comment":
    quick_prompt = "Write a report card comment. Ask me for the student's grade/performance level and any behavior notes if I haven't given them yet."
    st.session_state.mode = None  # reset immediately since no form needed

# ---------- Generate Notes form ----------
elif st.session_state.mode == "notes":
    with st.form("notes_form"):
        st.subheader("Generate Notes")
        subject = st.text_input("Subject", placeholder="e.g. Chemistry")
        topic = st.text_input("Topic / Sub-topic", placeholder="e.g. Periodic Table and Electron Configuration")
        grade = st.text_input("Class / Grade", placeholder="e.g. Form 1")
        depth = st.selectbox("Level of detail", ["Brief overview", "Standard (exam-ready)", "In-depth / detailed"])
        note_format = st.selectbox("Format", ["Bullet points", "Structured with headings and sub-points", "Numbered outline"])
        submitted = st.form_submit_button("Generate Notes")

        if submitted:
            quick_prompt = f"""Create clear, well-organized study notes for learners.

Subject: {subject}
Topic / Sub-topic: {topic}
Class / Grade level: {grade}
Level of detail: {depth}
Format: {note_format}

Structure the notes with:
- A short introduction explaining what the topic covers and why it matters
- Key definitions and terms clearly explained
- Main concepts broken into clearly labeled sections
- Simple examples or real-world applications where relevant
- A short summary/recap at the end (3-5 key takeaways)

Keep the language appropriate for the stated class/grade level. Use markdown headings and bullet points for readability."""

# ---------- Summarize Document form ----------
elif st.session_state.mode == "summarize":
    with st.form("summarize_form"):
        st.subheader("Summarize a Document")
        uploaded_file = st.file_uploader("Upload a .docx, .pdf, or .txt file", type=["docx", "pdf", "txt"])
        length = st.selectbox("Summary length", ["Short (a few sentences)", "Medium (a paragraph)", "Detailed (key points list)"])
        submitted = st.form_submit_button("Summarize")

        if submitted and uploaded_file is not None:
            file_text = read_uploaded_file(uploaded_file)
            if not file_text.strip():
                st.warning("Couldn't find any readable text in this file. If it's a scanned PDF (a photo of a page rather than typed text), text extraction won't work — try a typed/digital document instead.")
            else:
                # keep prompts fast — trim very long documents
                file_text = file_text[:12000]
                quick_prompt = f"""Summarize the following document for a busy teacher. Summary length: {length}.

Document:
\"\"\"
{file_text}
\"\"\""""
        elif submitted and uploaded_file is None:
            st.warning("Please upload a file first.")

# ---------- Example prompts (shown only when there's no conversation yet) ----------
if len(st.session_state.messages) == 0 and st.session_state.mode is None:
    import random
    example_prompts = [
        ("📝", "Make a lesson plan", "Create a lesson plan for a topic I'll describe"),
        ("❓", "Build a quick quiz", "Create a short quiz on a topic I'll describe"),
        ("💬", "Draft a report comment", "Write a report card comment for a student"),
        ("📓", "Summarize a concept", "Create simple study notes on a topic I'll describe"),
    ]
    random.shuffle(example_prompts)
    st.write("**Try one of these:**")
    ex_col1, ex_col2 = st.columns(2)
    for idx, (emoji, label, prompt_text) in enumerate(example_prompts):
        target_col = ex_col1 if idx % 2 == 0 else ex_col2
        with target_col:
            if st.button(f"{emoji} {label}", key=f"example_{idx}", use_container_width=True):
                quick_prompt = prompt_text

# ---------- Show past messages (with download/edit for assistant replies) ----------
for i, msg in enumerate(st.session_state.messages):
    avatar = "🧑‍🏫" if msg["role"] == "user" else "🤖"
    with st.chat_message(msg["role"], avatar=avatar):
        if msg["role"] == "assistant":
            edit_key = f"edit_mode_{i}"
            if edit_key not in st.session_state:
                st.session_state[edit_key] = False

            if st.session_state[edit_key]:
                edited_text = st.text_area("Edit this document:", value=msg["content"], height=300, key=f"textarea_{i}")
                if st.button("💾 Save changes", key=f"save_{i}"):
                    st.session_state.messages[i]["content"] = edited_text
                    if "id" in msg:
                        update_message(msg["id"], edited_text)
                    st.session_state[edit_key] = False
                    st.rerun()
            else:
                st.write(msg["content"])

            btn_col1, btn_col2, btn_col3 = st.columns(3)
            with btn_col1:
                if st.button("✏️ Edit", key=f"edit_btn_{i}", use_container_width=True):
                    st.session_state[edit_key] = True
                    st.rerun()
            with btn_col2:
                docx_buffer = text_to_docx_bytes(msg["content"])
                st.download_button(
                    "⬇️ Download",
                    data=docx_buffer,
                    file_name=f"document_{i}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_{i}",
                    use_container_width=True
                )
            with btn_col3:
                speak_button(msg["content"], key=f"speak_{i}")
        else:
            st.write(msg["content"])

# ---------- Custom styling for a rounded, Claude-style input bar, pinned to the bottom ----------
st.markdown("""
<style>
div[data-testid="stForm"] {
    border: 1px solid #d0d0d0;
    border-radius: 24px;
    padding: 8px 12px;
}
div[data-testid="stForm"] div[data-testid="stTextInput"] input {
    border: none !important;
    box-shadow: none !important;
}
div[data-testid="stForm"] button {
    border-radius: 20px !important;
    height: 42px;
}
.st-key-input_bar_wrapper button[kind="secondary"] {
    border-radius: 50% !important;
    height: 42px;
    width: 42px;
}
/* Stick the input bar to the bottom of the main content area (not the whole screen —
   this keeps it from overlapping the sidebar) */
.st-key-input_bar_wrapper {
    position: sticky;
    bottom: 10px;
    width: 100%;
    max-width: 700px;
    margin: 0 auto;
    border-radius: 26px;
    z-index: 999;
    padding: 12px 8px;
    border: 2px solid transparent;
    background:
        linear-gradient(rgba(255,255,255,0.9), rgba(255,255,255,0.9)) padding-box,
        linear-gradient(90deg, #EA4335, #4285F4, #34A853, #FBBC05, #EA4335) border-box;
    background-size: 100% 100%, 300% 100%;
    animation: glow-border-shift 6s linear infinite;
    box-shadow: 0 4px 24px rgba(66, 133, 244, 0.15), 0 4px 18px rgba(0,0,0,0.06);
}
@keyframes glow-border-shift {
    0% { background-position: 0 0, 0% 0; }
    100% { background-position: 0 0, 300% 0; }
}
/* Leave a little room at the bottom of the page so the last message isn't crowded */
.main .block-container {
    padding-bottom: 40px;
}
</style>
""", unsafe_allow_html=True)

# ---------- Voice recorder (appears when the mic button is tapped) ----------
if "show_recorder" not in st.session_state:
    st.session_state.show_recorder = False

if st.session_state.show_recorder:
    audio_value = st.audio_input("Tap to record, tap again to stop")
    if audio_value is not None:
        current_audio_id = getattr(audio_value, "file_id", audio_value.name)
        if st.session_state.get("last_audio_id") != current_audio_id:
            with st.spinner("Transcribing your voice..."):
                try:
                    transcription = client.audio.transcriptions.create(
                        file=("recording.wav", audio_value.getvalue()),
                        model="whisper-large-v3-turbo"
                    )
                    st.session_state.last_audio_id = current_audio_id
                    st.session_state.pending_voice_prompt = transcription.text
                    st.session_state.show_recorder = False
                    st.rerun()
                except Exception as e:
                    st.error(f"Couldn't transcribe that: {e}")

# ---------- Input row: mic button (outside the form) + text box + send button (inside the form) ----------
user_input = None

with st.container(key="input_bar_wrapper"):
    col_mic, col_form = st.columns([1, 9])
    with col_mic:
        if st.button("🎤", key="mic_toggle_btn"):
            st.session_state.show_recorder = not st.session_state.show_recorder
            st.rerun()

    with col_form:
        with st.form("chat_input_form", clear_on_submit=True):
            col_text, col_send = st.columns([8, 1])
            with col_text:
                typed_text = st.text_input(
                    "message", placeholder="Write a message...", label_visibility="collapsed"
                )
            with col_send:
                send_clicked = st.form_submit_button("➤")

if send_clicked and typed_text.strip():
    user_input = typed_text.strip()

# If a form was just submitted, use that prompt instead
if quick_prompt:
    user_input = quick_prompt
    st.session_state.mode = None  # close the form after submitting

# If a voice recording was just transcribed, use that instead
if st.session_state.get("pending_voice_prompt"):
    user_input = st.session_state.pending_voice_prompt
    st.session_state.pending_voice_prompt = None

if user_input:
    # If this is the first message in the conversation, auto-title it
    if len(st.session_state.messages) == 0:
        auto_title = user_input.strip()[:40] + ("..." if len(user_input.strip()) > 40 else "")
        rename_conversation(st.session_state.current_conversation_id, auto_title)

    # Show the user's message and save it to the database
    user_msg_id = save_message(st.session_state.current_conversation_id, "user", user_input)
    st.session_state.messages.append({"id": user_msg_id, "role": "user", "content": user_input})
    with st.chat_message("user", avatar="🧑‍🏫"):
        st.write(user_input)

    # Send it to Groq and show the reply
    with st.chat_message("assistant", avatar="🤖"):
        with st.spinner("Thinking..."):
            try:
                response = client.chat.completions.create(
                    model="openai/gpt-oss-120b",
                    max_tokens=2000,
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant for teachers. You help create lesson plans, tests, quizzes, assignments, grading rubrics, report card comments, and document summaries. Follow any structural instructions given exactly and precisely. Be practical, clear, and well-formatted using markdown headings and tables where appropriate."},
                        {"role": "user", "content": user_input}
                    ]
                )
                reply = response.choices[0].message.content
                st.write(reply)
            except Exception as e:
                reply = f"Sorry, something went wrong: {e}"
                st.error(reply)

    # Save the reply to the database and history
    assistant_msg_id = save_message(st.session_state.current_conversation_id, "assistant", reply)
    st.session_state.messages.append({"id": assistant_msg_id, "role": "assistant", "content": reply})
    st.rerun()